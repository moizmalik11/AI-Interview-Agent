from datetime import datetime
import json
import csv
import os
import joblib
import numpy as np
from collections import Counter, defaultdict
import pandas as pd
import numpy as np
from docx import Document
from utils.model import make_prediction
from docx.shared import Inches
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

def safe_entropy(probs):
    return -np.sum(probs * np.log2(probs + 1e-9))

def extract_eye_tracking(data, duration):
    blink_timestamps = []
    eye_directions = []
    streak = 0
    max_streak = 0
    for frame in data:
        eye_info = frame.get("eye_tracking", {})
        ts = frame.get("timestamp", 0)
        eye_directions.append(eye_info.get("direction", "center"))
        if eye_info.get("blink_detected", False):
            blink_timestamps.append(ts)
            streak += 1
        else:
            max_streak = max(max_streak, streak)
            streak = 0

    early, mid, late = 0, 0, 0
    for ts in blink_timestamps:
        if ts < duration / 3:
            early += 1
        elif ts < 2 * duration / 3:
            mid += 1
        else:
            late += 1

    eye_dir_counts = Counter(eye_directions)
    total = len(eye_directions)
    eye_dist = {f"eye_dir_{k}": eye_dir_counts.get(k, 0) / total for k in ["left", "right", "center", "up", "down"]}
    entropy = safe_entropy(np.array(list(eye_dist.values())))

    return {
        "total_blinks": len(blink_timestamps),
        "blink_rate_per_sec": len(blink_timestamps) / duration,
        "blinks_early": early,
        "blinks_mid": mid,
        "blinks_late": late,
        "max_blink_streak": max_streak,
        "eye_direction_entropy": entropy,
        **eye_dist
    }

def extract_head_tracking(data, duration):
    yaws, pitches, rolls = [], [], []
    rotation_deltas = []
    jerks = 0
    last = None
    for frame in data:
        head = frame.get("head_tracking", {})
        yaw, pitch, roll = head.get("yaw", 0), head.get("pitch", 0), head.get("roll", 0)
        yaws.append(yaw)
        pitches.append(pitch)
        rolls.append(roll)
        if last:
            delta = np.linalg.norm([yaw - last[0], pitch - last[1], roll - last[2]])
            rotation_deltas.append(delta)
            if delta > 15:
                jerks += 1
        last = (yaw, pitch, roll)
    return {
        "avg_yaw": np.mean(yaws),
        "avg_pitch": np.mean(pitches),
        "avg_roll": np.mean(rolls),
        "head_rotation_variance": np.var(rotation_deltas),
        "head_jerk_rate": jerks / duration,
        "head_rotation_per_sec": np.sum(rotation_deltas) / duration
    }

def extract_emotion(data, duration):
    emotions = []
    emotion_time = defaultdict(float)
    switches = 0
    last_emotion = None
    first_switch_time = None
    for frame in data:
        e = frame.get("emotion", {}).get("detected_emotion", "neutral")
        emotions.append(e)
        emotion_time[e] += 1
        if last_emotion and e != last_emotion:
            switches += 1
            if first_switch_time is None:
                first_switch_time = frame.get("timestamp", duration)
        last_emotion = e

    dom = Counter(emotions).most_common(1)[0][0] if emotions else "neutral"
    emotion_distr = {f"emotion_{k}": v / duration for k, v in emotion_time.items()}
    return {
        "dominant_emotion": dom,
        "emotion_switches": switches,
        "emotion_switch_rate": switches / duration,
        "time_to_first_emotion_switch": first_switch_time or duration,
        "final_emotion": emotions[-1] if emotions else "neutral",
        **emotion_distr
    }

def extract_body_movement(data):
    movements = 0
    burst_count = 0
    streak = 0
    for frame in data:
        move = frame.get("body_tracking", {}).get("movement", "No Movement")
        if move != "No Movement":
            movements += 1
            streak += 1
        else:
            if streak >= 2:
                burst_count += 1
            streak = 0
    return {
        "body_movement_rate": movements / len(data),
        "body_movement_bursts": burst_count
    }

def extract_object_detection(data):
    obj_set = set()
    total_objs = 0
    person_presence = 0
    presence_gaps = 0
    last_present = True
    suspicious_labels = {"laptop", "phone", "monitor", "tv", "screen", "keyboard", "mouse", "tablet", "camera"}
    suspicious_flag = False
    for frame in data:
        labels = [o["label"].lower() for o in frame.get("detections", {}).get("yolo", [])]
        total_objs += len(labels)
        obj_set.update(labels)
        if any(label in suspicious_labels for label in labels):
            suspicious_flag = True
        person_here = 'person' in labels
        if not person_here and last_present:
            presence_gaps += 1
        if person_here:
            person_presence += 1
        person_count = labels.count("person")
        if person_count > 1:
            suspicious_flag = True  # mark as suspicious

        last_present = person_here
    return {
        "avg_objects_per_frame": total_objs / len(data),
        "unique_objects_count": len(obj_set),
        "person_visible_ratio": person_presence / len(data),
        "person_presence_gaps": presence_gaps,
        "suspicious_object_detected": int(suspicious_flag)
    }

def extract_engagement_score(eye, head, obj):
    score = (
        (1 - head["head_rotation_variance"] / 1000) * 0.3 +
        eye.get("eye_dir_center", 0) * 0.4 +
        obj["person_visible_ratio"] * 0.3
    )
    return max(0, min(1, score))

def extract_features(json_path):
    with open(json_path, 'r') as f:
        data = json.load(f)
    if not data:
        return None
    duration = data[-1].get("timestamp", 0) - data[0].get("timestamp", 0)
    duration = max(duration, 1.0)
    eye = extract_eye_tracking(data, duration)
    head = extract_head_tracking(data, duration)
    emo = extract_emotion(data, duration)
    body = extract_body_movement(data)
    obj = extract_object_detection(data)
    engage = extract_engagement_score(eye, head, obj)
    return {
        "filename": os.path.basename(json_path),
        "duration": duration,
        "engagement_score": round(engage, 3),
        **eye, **head, **emo, **body, **obj
    }

def generate_csv_rowwise(user_dir,date_str ,user_name):
    f_data = extract_features(f"{user_dir}/{date_str}.json")
    result = make_prediction(f_data)
    print(f"🔍 make_prediction {result}")

    if not f_data:
        print("❌ Feature extraction failed.")
        return

    f_data["name"] = user_name
    f_data["filename"] = f"{user_dir}/{date_str}"
    f_data["duration"] = round(f_data["duration"] / 60, 2)  # convert seconds to minutes
    f_data["engagement_score"] = round(f_data["engagement_score"], 2)
    label, score, reason = assign_label(f_data)
    f_data["label"] = label
    f_data["score"] = score
    f_data["reason"] = reason

    print(f'lable: {f_data["label"]}')

    with open(f"{user_dir}/{date_str}.csv", 'w', newline='') as f:
        writer = csv.writer(f)
        writer.writerow(["Feature", "Value"])
        for key, value in f_data.items():
            writer.writerow([key, value])

    print(f"✅ Row-wise CSV saved at {user_dir}/{date_str}.csv")
    generate_word_report(f_data, f"{user_dir}/{date_str}.docx")


    

def assign_label(features):
    debug_info = {}
    score = 0

    if features.get("suspicious_object_detected", 0) == 1:
        person_changed = features.get("person_changed", 0) == 1
        person_count = features.get("person_count", 1)
        if person_changed or person_count > 1:
            debug_info["forced_cheating"] = (
                "Suspicious object + person change/multiple persons"
            )
            return "cheating", 10.0, debug_info

    weight = {
        "suspicious_object": 1.5,     
        "presence_gap": 1.0,
        "low_visibility": 1.0,
        "head_jerk": 1.0,
        "abnormal_blink": 0.8,
        "emotion_switch": 0.8,
        "low_engagement": 1.5,
    }

    if features.get("suspicious_object_detected", 0) == 1:
        score += weight["suspicious_object"]
        debug_info["suspicious_object"] = "Suspicious object detected"

    if features.get("person_presence_gaps", 0) >= 1:
        score += weight["presence_gap"]
        debug_info["presence_gap"] = f"Presence gaps: {features.get('person_presence_gaps')}"

    if features.get("person_visible_ratio", 1.0) < 0.8:
        score += weight["low_visibility"]
        debug_info["low_visibility"] = f"Visible ratio too low: {features.get('person_visible_ratio')}"

    if features.get("head_jerk_rate", 0.0) > 0.5:
        score += weight["head_jerk"]
        debug_info["head_jerk"] = f"Head jerk rate high: {features.get('head_jerk_rate')}"

    blink_rate = features.get("blink_rate_per_sec", 0.3)
    if blink_rate < 0.1 or blink_rate > 0.6:
        score += weight["abnormal_blink"]
        debug_info["abnormal_blink"] = f"Abnormal blink rate: {blink_rate}"

    if features.get("emotion_switch_rate", 0.0) > 0.5:
        score += weight["emotion_switch"]
        debug_info["emotion_switch"] = f"Emotion switch rate high: {features.get('emotion_switch_rate')}"

    if features.get("engagement_score", 1.0) < 0.4:
        score += weight["low_engagement"]
        debug_info["low_engagement"] = f"Low engagement: {features['engagement_score']}"


    label = "cheating" if score >= 3.5 else "not_cheating"
    return label, round(score, 2), debug_info




def generate_word_report(features, output_path):
    doc = Document()
    doc.add_heading('Cheating Detection Report', 0)

    # Table for basic details
    doc.add_heading('Candidate Details', level=1)
    details_table = doc.add_table(rows=0, cols=2)
    details_table.style = 'Table Grid'

    fields = ['filename', 'name', 'duration', 'engagement_score', 'label']
    for field in fields:
        row_cells = details_table.add_row().cells
        row_cells[0].text = field.replace("_", " ").title()
        row_cells[1].text = str(features.get(field, ""))

    doc.add_paragraph()

    reasons = features.get("reason")
    if isinstance(reasons, dict):
        doc.add_heading("Reasons for Classification", level=1)
        reasons_table = doc.add_table(rows=1, cols=2)
        reasons_table.style = 'Light Grid Accent 1'

        hdr_cells = reasons_table.rows[0].cells
        hdr_cells[0].text = 'Reason'
        hdr_cells[1].text = 'Explanation'

        for reason, explanation in reasons.items():
            row_cells = reasons_table.add_row().cells
            row_cells[0].text = str(reason)
            row_cells[1].text = str(explanation)

        doc.add_paragraph()

    doc.add_heading('Feature Breakdown', level=1)
    for key, value in features.items():
        if key not in set(fields + ['reason']):
            doc.add_paragraph(f"{key}: {value}", style='List Bullet')

    doc.save(output_path)
    print(f"✅ Word report saved at {output_path}")


def predict_cheating(user_dir,date_str, user_name):
    generate_csv_rowwise(user_dir,date_str,user_name)

